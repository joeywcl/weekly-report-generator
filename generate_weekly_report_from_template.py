import argparse
from pathlib import Path
import yaml
from datetime import datetime, timedelta
from typing import Optional
from docx import Document
from docx.shared import Inches
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement


def get_week_range(today: Optional[datetime] = None) -> str:
    """Calculate week range (Monday to Friday) for the given date (defaults to today)."""
    today = today or datetime.now()
    monday = today - timedelta(days=today.weekday())
    friday = monday + timedelta(days=4)
    return f"{monday.strftime('%Y-%m-%d')} → {friday.strftime('%Y-%m-%d')}"


def delete_paragraph(paragraph: Paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None


def insert_paragraph_after(paragraph: Paragraph, text: str, style: str = None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text is not None:
        new_para.add_run(text)
    return new_para


def insert_paragraph_after_runs(paragraph: Paragraph, runs: list, style: str = None):
    """Insert a paragraph with multiple runs; each run is (text, bold)."""
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    for text, bold in runs:
        if text:
            r = new_para.add_run(text)
            r.bold = bold
    return new_para


def _set_left_indent(paragraph: Paragraph, left_indent_in: Optional[float]):
    if left_indent_in is None:
        return
    paragraph.paragraph_format.left_indent = Inches(left_indent_in)


def insert_paragraph_after_i(
    paragraph: Paragraph,
    text: str,
    style: str = None,
    left_indent_in: Optional[float] = None,
):
    p = insert_paragraph_after(paragraph, text, style=style)
    _set_left_indent(p, left_indent_in)
    return p


def insert_paragraph_after_runs_i(
    paragraph: Paragraph,
    runs: list,
    style: str = None,
    left_indent_in: Optional[float] = None,
):
    p = insert_paragraph_after_runs(paragraph, runs, style=style)
    _set_left_indent(p, left_indent_in)
    return p


def insert_label_value(
    anchor: Paragraph,
    label: str,
    value: str,
    style: str,
    left_indent_in: Optional[float] = None,
):
    """Insert a single line: bold label + normal value."""
    return insert_paragraph_after_runs_i(
        anchor,
        [(label, True), (value, False)],
        style=style,
        left_indent_in=left_indent_in,
    )


def insert_label_value_block(
    anchor: Paragraph,
    label: str,
    value: str,
    style: str,
    left_indent_in: Optional[float] = None,
):
    """Insert label/value on same line when simple, otherwise label then parsed content."""
    v = (value or "").strip("\n")
    if not v.strip():
        return insert_label_value(
            anchor, label, " N/A", style=style, left_indent_in=left_indent_in
        )
    simple_one_liner = ("\n" not in v) and (not v.lstrip().startswith("-"))
    if simple_one_liner:
        return insert_label_value(
            anchor, label, (" " + v) if v else "", style, left_indent_in
        )
    last = insert_paragraph_after_runs_i(
        anchor, [(label, True)], style=style, left_indent_in=left_indent_in
    )
    return insert_content_parsed(last, v, style=style, left_indent_in=left_indent_in)


def insert_content_parsed(
    anchor: Paragraph,
    content: str,
    style: str = "Normal",
    left_indent_in: Optional[float] = None,
):
    """Insert content where lines starting with '-' are bullets, others are normal text.

    left_indent_in applies a visual left indent for all inserted paragraphs.
    Returns last paragraph.
    """
    c = (content or "").replace("\r\n", "\n").strip("\n")
    if not c.strip():
        return insert_paragraph_after_i(
            anchor, "N/A", style=style, left_indent_in=left_indent_in
        )
    last = anchor
    for line in c.splitlines():
        line = line.rstrip()
        if not line:
            continue
        stripped = line.lstrip()
        if stripped.startswith("-"):
            bullet_text = stripped[1:].strip()
            last = insert_paragraph_after_i(
                last, "• " + bullet_text, style=style, left_indent_in=left_indent_in
            )
        else:
            last = insert_paragraph_after_i(
                last, line.strip(), style=style, left_indent_in=left_indent_in
            )
    return last


def find_paragraph(doc: Document, exact_text: str):
    for p in doc.paragraphs:
        if p.text.strip() == exact_text.strip():
            return p
    return None


def set_line_value(doc: Document, prefix: str, value: str):
    # Finds "Prefix: ..." lines like "Name: X" and replaces only the text.
    for p in doc.paragraphs:
        if p.text.strip().startswith(prefix):
            p.text = f"{prefix}{value}"
            return
    raise ValueError(f"Could not find line starting with '{prefix}'")


def clear_between(
    doc: Document, start_heading_text: str, stop_heading_texts: list[str]
):
    start_p = find_paragraph(doc, start_heading_text)
    if not start_p:
        raise ValueError(f"Start heading not found: {start_heading_text}")
    start_el = start_p._element

    started = False
    to_delete = []
    for p in doc.paragraphs:
        if p._element == start_el:
            started = True
            continue
        if not started:
            continue
        if p.text.strip() in stop_heading_texts:
            break
        to_delete.append(p)

    for p in to_delete:
        delete_paragraph(p)

    # re-find and return the heading paragraph (safe after deletions)
    start_p2 = find_paragraph(doc, start_heading_text)
    if not start_p2:
        raise ValueError(f"Start heading disappeared: {start_heading_text}")
    return start_p2


def main():
    ap = argparse.ArgumentParser(
        description="Generate weekly report docx from your template with consistent formatting."
    )
    ap.add_argument(
        "--template",
        required=True,
        help="Path to the Word template .docx (your fixed format).",
    )
    ap.add_argument("--input", required=True, help="Path to YAML input file.")
    ap.add_argument("--output", required=True, help="Path to output .docx.")
    args = ap.parse_args()

    data = yaml.safe_load(Path(args.input).read_text(encoding="utf-8"))

    doc = Document(args.template)

    # Header lines
    set_line_value(doc, "Name: ", data.get("name", ""))
    set_line_value(doc, "Role: ", data.get("role", ""))
    week_value = data.get("week")
    if not week_value or not str(week_value).strip():
        week_value = get_week_range()
    set_line_value(doc, "Week: ", str(week_value))

    # Weekly objective: replace the paragraph(s) after the objective heading, up to next Heading 3.
    obj_anchor = clear_between(
        doc,
        "Weekly Objective (One Sentence)",
        stop_heading_texts=["Execution & Output", "The “2X” Transformation Log"],
    )
    # Insert one objective line in same style as template used (List Paragraph)
    weekly_objective = (data.get("weekly_objective", "") or "").strip()
    if not weekly_objective:
        weekly_objective = "N/A"
    obj_line = insert_paragraph_after(
        obj_anchor, weekly_objective, style="List Paragraph"
    )
    # Spacer before next heading
    insert_paragraph_after(obj_line, "", style="List Paragraph")

    # Execution & Output: replace list items until next heading
    exe_anchor = clear_between(
        doc, "Execution & Output", stop_heading_texts=["The “2X” Transformation Log"]
    )
    # Execution: each item = bold summary, then content (lines starting with "-" = bullet, else text)
    exec_items = data.get("execution_output", []) or []
    last = exe_anchor
    for idx, item in enumerate(exec_items):
        if isinstance(item, dict):
            summary = item.get("summary", "")
            content = item.get("content", "")
        else:
            summary = ""
            content = str(item) if item else ""
        # Legacy: description + bullets + order
        if not content and isinstance(item, dict):
            description = item.get("description", "")
            bullets = item.get("bullets", []) or []
            if isinstance(bullets, str):
                bullets = [b.strip() for b in bullets.splitlines() if b.strip()]
            order = item.get("order", "bullets_first")
            parts = []
            if order == "description_first":
                if description:
                    parts.append(description)
                for b in bullets:
                    parts.append("- " + b)
            else:
                for b in bullets:
                    parts.append("- " + b)
                if description:
                    parts.append(description)
            content = "\n".join(parts)
        if summary or content:
            if summary:
                last = insert_paragraph_after_runs(
                    last, [(summary, True)], style="List Paragraph"
                )
            c = (content or "").replace("\r\n", "\n").strip("\n")
            if not c.strip():
                last = insert_paragraph_after(last, "N/A", style="List Paragraph")
            for line in c.splitlines():
                line = line.rstrip()
                if not line:
                    continue
                stripped = line.lstrip()
                if stripped.startswith("-"):
                    bullet_text = stripped[1:].strip()
                    last = insert_paragraph_after(
                        last, "• " + bullet_text, style="List Paragraph"
                    )
                else:
                    last = insert_paragraph_after(
                        last, line.strip(), style="List Paragraph"
                    )
            if idx < len(exec_items) - 1:
                last = insert_paragraph_after(last, "", style="List Paragraph")

    # Spacer before next heading
    insert_paragraph_after(last, "", style="List Paragraph")

    # AI Acceleration: replace paragraphs after "AI Acceleration" until "SOP & Process Solidification"
    ai_anchor = clear_between(
        doc,
        "AI Acceleration",
        stop_heading_texts=["SOP & Process Solidification", "Friction, Blockers & Ask"],
    )
    tasks = (data.get("transformation_log", {}) or {}).get(
        "ai_acceleration_tasks", []
    ) or []
    last = ai_anchor
    for idx, t in enumerate(tasks):
        indent_in = 0.25
        style = "First Paragraph" if idx == 0 else "Body Text"
        last = insert_label_value(
            last,
            f"Task {idx + 1}: ",
            t.get("task", ""),
            style=style,
            left_indent_in=indent_in,
        )
        last = insert_label_value(
            last,
            "Tool / Agent: ",
            t.get("tool_agent", ""),
            style=style,
            left_indent_in=indent_in,
        )
        last = insert_label_value(
            last,
            "Time Saved (Est.): ",
            t.get("time_saved", ""),
            style=style,
            left_indent_in=indent_in,
        )
        last = insert_paragraph_after_runs_i(
            last,
            [("Insight / Limitation:", True)],
            style=style,
            left_indent_in=indent_in,
        )
        last = insert_content_parsed(
            last,
            t.get("insight_failure", ""),
            style=style,
            left_indent_in=indent_in,
        )
        if idx < len(tasks) - 1:
            last = insert_paragraph_after(last, "", style=style)

    # Spacer before next heading
    insert_paragraph_after(last, "", style="Body Text")

    # SOP & Process Solidification: list of Item/Impact (or legacy single item/impact)
    sop_data = data.get("sop_process_solidification", {}) or {}
    sop_items = sop_data.get("items", [])
    if not sop_items and ("item" in sop_data or "impact" in sop_data):
        sop_items = [
            {
                "item": sop_data.get("item", "None"),
                "impact": sop_data.get("impact", "N/A"),
            }
        ]
    if not sop_items:
        sop_items = [{"item": "None", "impact": "N/A"}]
    sop_anchor = find_paragraph(doc, "SOP & Process Solidification")
    if sop_anchor:
        # clear until friction heading
        clear_between(
            doc,
            "SOP & Process Solidification",
            stop_heading_texts=[
                "Friction, Blockers & Ask",
                "Next Week’s Focus (Preview Only) ",
                "Next Week’s Focus (Preview Only)",
            ],
        )
        last = sop_anchor
        for idx, s in enumerate(sop_items):
            indent_in = 0.25
            item = s.get("item", "None") if isinstance(s, dict) else "None"
            impact = s.get("impact", "N/A") if isinstance(s, dict) else "N/A"
            last = insert_label_value(
                last,
                "Item: ",
                item,
                style="Normal",
                left_indent_in=indent_in,
            )
            last = insert_label_value_block(
                last,
                "Impact:",
                impact,
                style="Normal",
                left_indent_in=indent_in,
            )
            if idx < len(sop_items) - 1:
                last = insert_paragraph_after(last, "", style="Normal")

        # Spacer before next heading
        insert_paragraph_after(last, "", style="Normal")

    # Friction section: replace paragraphs until Next Week’s Focus heading
    fr_anchor = clear_between(
        doc,
        "Friction, Blockers & Ask",
        stop_heading_texts=[
            "Next Week’s Focus (Preview Only) ",
            "Next Week’s Focus (Preview Only)",
        ],
    )
    # Friction 1, Friction 2, ... each with Friction description:, Action/Mitigation:, Ask/Attention needed: (bold labels)
    fr_items = data.get("friction_blockers_ask", []) or []
    last = fr_anchor
    for idx, item in enumerate(fr_items, 1):
        indent_in = 0.25
        if isinstance(item, dict):
            friction = item.get("friction", "")
            action = item.get("action_mitigation", "")
            ask = item.get("ask_attention_needed", "")
        else:
            friction = str(item) if item else ""
            action = ""
            ask = ""
        if friction or action or ask:
            # Keep friction title + description on one line (match Focus/Task style).
            fr_value = (friction or "").strip() or "N/A"
            last = insert_label_value(
                last,
                f"Friction {idx}: ",
                fr_value,
                style="Normal",
                left_indent_in=indent_in,
            )
            last = insert_label_value_block(
                last,
                "Action/Mitigation:",
                action,
                style="Normal",
                left_indent_in=indent_in,
            )
            last = insert_label_value_block(
                last,
                "Ask/Attention needed:",
                ask,
                style="Normal",
                left_indent_in=indent_in,
            )
            if idx < len(fr_items):
                last = insert_paragraph_after(last, "", style="Normal")
    insert_paragraph_after(last, "", style="Normal")

    # Next Week focus: replace until end
    nw_heading = "Next Week’s Focus (Preview Only) "
    nw_anchor = find_paragraph(doc, nw_heading)
    if not nw_anchor:
        nw_heading = "Next Week’s Focus (Preview Only)"
        nw_anchor = find_paragraph(doc, nw_heading)
    if nw_anchor:
        # delete everything after this heading
        started = False
        to_delete = []
        for p in doc.paragraphs:
            if p._element == nw_anchor._element:
                started = True
                continue
            if started:
                to_delete.append(p)
        for p in to_delete:
            delete_paragraph(p)

        focus_lines = data.get("next_week_focus", []) or []
        last = nw_anchor
        for i, line in enumerate(focus_lines, 1):
            last = insert_label_value(
                last,
                f"Focus {i}: ",
                line,
                style="Normal",
                left_indent_in=0.25,
            )

    Path(args.output).parent.mkdir(parents=True, exist_ok=True)
    doc.save(args.output)


if __name__ == "__main__":
    main()
